"""
generate_resume_ar.py
Generates Ayman_AlShehri_Resume_AR.docx — Arabic one-page recruiter resume (RTL).

Usage:
    pip install python-docx
    python generate_resume_ar.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Colors ───────────────────────────────────────────────────────────────────
TEAL  = RGBColor(0x0d, 0x94, 0x88)
NAVY  = RGBColor(0x0f, 0x17, 0x2a)
MUTED = RGBColor(0x64, 0x74, 0x8b)
BODY  = RGBColor(0x1e, 0x29, 0x3b)

# ── Layout ───────────────────────────────────────────────────────────────────
MARGIN        = 0.6
PAGE_WIDTH    = 8.5
CONTENT_WIDTH = PAGE_WIDTH - 2 * MARGIN

ARABIC_FONT = 'Arial'


# ── RTL helpers ───────────────────────────────────────────────────────────────

def rtl_paragraph(paragraph):
    """Mark a paragraph as RTL and right-align it."""
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:bidi')):
        pPr.remove(old)
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def font(run, size, bold=False, color=None, arabic=True):
    """
    Apply font. When arabic=True, sets complex-script properties so Word
    renders Arabic glyphs correctly in RTL:
      - w:rFonts w:cs  (complex script font)
      - w:szCs         (complex script size)
      - w:bCs          (complex script bold)
      - w:lang w:bidi  (language: ar-SA)
      - w:rtl          (RTL run direction)
    """
    run.font.name = ARABIC_FONT
    run.font.size = Pt(size)
    run.bold      = bold
    if color:
        run.font.color.rgb = color

    if arabic:
        rPr = run._r.get_or_add_rPr()

        # Complex script font name
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), ARABIC_FONT)
        rFonts.set(qn('w:hAnsi'), ARABIC_FONT)
        rFonts.set(qn('w:cs'),    ARABIC_FONT)

        # Complex script font size (half-points)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(size * 2)))
        rPr.append(szCs)

        # Complex script bold
        if bold:
            rPr.append(OxmlElement('w:bCs'))

        # Language: Arabic (Saudi Arabia)
        lang = rPr.find(qn('w:lang'))
        if lang is None:
            lang = OxmlElement('w:lang')
            rPr.append(lang)
        lang.set(qn('w:bidi'), 'ar-SA')

        # RTL run direction
        if rPr.find(qn('w:rtl')) is None:
            rPr.append(OxmlElement('w:rtl'))


def spacing(paragraph, before=0, after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing = Pt(13)


def bottom_border(paragraph, color='0d9488', sz='4'):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Document builders ─────────────────────────────────────────────────────────

def section_header(doc, text):
    p = doc.add_paragraph()
    spacing(p, before=5, after=1)
    rtl_paragraph(p)
    r = p.add_run(text)
    font(r, 10, bold=True, color=NAVY)
    bottom_border(p)
    return p


def company_line(doc, company_loc, dates, top=3):
    """
    Single RTL paragraph: company on right, dates trailing to the left.
    In an RTL paragraph the logically-first text appears on the RIGHT and
    the logically-last text ends up on the LEFT — exactly what we want.
    """
    p = doc.add_paragraph()
    spacing(p, before=top, after=0)
    rtl_paragraph(p)
    r = p.add_run(company_loc)
    font(r, 10, bold=True, color=NAVY)
    r = p.add_run('   \u2014   ' + dates)
    font(r, 9, color=MUTED)
    return p


def role_line(doc, role):
    """Role title in teal. No italic — Arabic script doesn't render italic cleanly."""
    p = doc.add_paragraph()
    spacing(p, before=0, after=1)
    rtl_paragraph(p)
    r = p.add_run(role)
    font(r, 9.5, color=TEAL)
    return p


def bullet(doc, text):
    p = doc.add_paragraph()
    spacing(p, before=1, after=0)
    rtl_paragraph(p)
    pf = p.paragraph_format
    pf.right_indent = Inches(0.15)   # indent block from right margin
    r = p.add_run('\u2022 ')          # bullet char
    font(r, 9.5, color=MUTED)
    r = p.add_run(text)
    font(r, 9.5, color=BODY)
    return p


def skill_row(doc, label, value):
    p = doc.add_paragraph()
    spacing(p, before=2, after=0)
    rtl_paragraph(p)
    r = p.add_run(label + ':  ')
    font(r, 9.5, bold=True, color=NAVY)
    r = p.add_run(value)
    font(r, 9.5, color=BODY)
    return p


def edu_row(doc, title, detail):
    p = doc.add_paragraph()
    spacing(p, before=2, after=0)
    rtl_paragraph(p)
    r = p.add_run(title + '  ')
    font(r, 9.5, bold=True, color=NAVY)
    r = p.add_run(detail)
    font(r, 9.5, color=MUTED)
    return p


# ── Main ──────────────────────────────────────────────────────────────────────

def create_resume_ar():
    doc = Document()

    # Page layout
    for sec in doc.sections:
        sec.page_width    = Inches(PAGE_WIDTH)
        sec.page_height   = Inches(11)
        sec.left_margin   = Inches(MARGIN)
        sec.right_margin  = Inches(MARGIN)
        sec.top_margin    = Inches(MARGIN)
        sec.bottom_margin = Inches(MARGIN)

    # Normal style: set font + RTL as document-wide default
    normal = doc.styles['Normal']
    normal.font.name = ARABIC_FONT
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(0)
    pPr = normal._element.get_or_add_pPr()
    pPr.append(OxmlElement('w:bidi'))   # RTL default for all paragraphs

    # ── HEADER ────────────────────────────────────────────────────────────────

    p = doc.add_paragraph()
    spacing(p, before=0, after=1)
    rtl_paragraph(p)
    r = p.add_run('أيمن الشهري')
    font(r, 22, bold=True, color=NAVY)

    p = doc.add_paragraph()
    spacing(p, before=0, after=2)
    rtl_paragraph(p)
    r = p.add_run('مدير تقنية المعلومات وقائد التميز الرقمي  |  PMI-ACP  |  CKAD')
    font(r, 10, color=TEAL)

    # Contact line: mostly Latin — arabic=False, but paragraph is still RTL
    p = doc.add_paragraph()
    spacing(p, before=0, after=4)
    rtl_paragraph(p)
    r = p.add_run(
        'aymalsh@gmail.com  |  linkedin.com/in/aymanshe  |  '
        'aymanshehri.me  |  github.com/AymanShe  |  '
        '+966 550 832 888  |  '
    )
    font(r, 9, color=MUTED, arabic=False)
    r = p.add_run('الرياض، المملكة العربية السعودية')
    font(r, 9, color=MUTED)
    bottom_border(p, color='0d9488', sz='6')

    # ── SUMMARY ───────────────────────────────────────────────────────────────

    section_header(doc, 'الملخص')

    p = doc.add_paragraph()
    spacing(p, before=3, after=0)
    rtl_paragraph(p)
    r = p.add_run(
        'مدير تقنية معلومات يمتلك أكثر من 7 سنوات من الخبرة في هندسة البرمجيات والقيادة التقنية '
        'وحوكمة المؤسسات. يقود حالياً البنية التشغيلية في RQM Business Solutions — أطر GRC '
        'وتوحيد عمليات تقنية المعلومات وحوكمة تسليم المشاريع. '
        'حاصل على شهادتَي PMI-ACP وCKAD؛ يتمتع بمصداقية متساوية '
        'في جلسات السياسات ومراجعات الهندسة المعمارية.'
    )
    font(r, 9.5, color=BODY)

    # ── EXPERIENCE ────────────────────────────────────────────────────────────

    section_header(doc, 'الخبرة العملية')

    # RQM
    company_line(doc, 'RQM Business Solutions \u2014 الرياض، المملكة العربية السعودية',
                 'أكتوبر 2024 \u2013 الحاضر')
    role_line(doc, 'مدير التميز الرقمي (أكتوبر 2025 \u2013 الحاضر)  /  قائد تقني (أكتوبر 2024 \u2013 أكتوبر 2025)')
    bullet(doc,
        'صمّم وطبّق إطار GRC من الصفر؛ ووحّد عمليات تقنية المعلومات شاملاً إدارة الأجهزة '
        'والتحكم في الوصول واشتراكات SaaS وعمليات الانضمام والمغادرة.')
    bullet(doc,
        'أرسى حوكمة Sprint وبروتوكولات التحكم في التغيير — وتفاوض على تغييرات النطاق باحترافية '
        'مع أصحاب المصلحة والعملاء، مع الحفاظ على طاقة الفريق والتزامات التسليم.')
    bullet(doc,
        'امتلك الهندسة المعمارية وخارطة الطريق التقنية ومعايير الجودة لمنصة SaaS السحابية '
        'عبر فريق من 4\u20136 مهندسين.')

    # Canada College
    company_line(doc, 'Canada College \u2014 مونتريال، كيبيك',
                 'يوليو 2022 \u2013 مايو 2025 (دوام جزئي)')
    role_line(doc, 'مدرّس تقني ومستشار التدريب الميداني')
    bullet(doc,
        'صمّم وقدّم منهجاً كاملاً في تصميم قواعد البيانات بنسبة نجاح 100%؛ '
        'وأدار أكثر من 50 تدريباً ميدانياً مع أصحاب العمل في قطاع التقنية بمونتريال.')

    # Genetec
    company_line(doc, 'Genetec \u2014 مونتريال، كيبيك', 'سبتمبر 2022 \u2013 أبريل 2023')
    role_line(doc, 'مطور برمجيات')
    bullet(doc,
        'أطلق ميزة مشغّل الوسائط من التصميم حتى الإنتاج على منصة Clearance '
        'لإدارة الأدلة الرقمية؛ وحلّ خللاً في سلامة البيانات عبر Microservices '
        'كان يُفسد ظهور ملفات القضايا للمحققين.')

    # Ministry of Justice
    company_line(doc, 'وزارة العدل \u2014 الرياض، المملكة العربية السعودية',
                 'يناير 2019 \u2013 ديسمبر 2021')
    role_line(doc, 'مطور برمجيات')
    bullet(doc,
        'امتلك تصميم وتسليم وحدة القضايا الاجتماعية ضمن منصة ناجز الوطنية '
        'التي تخدم ملايين المواطنين السعوديين.')
    bullet(doc,
        'قاد فريقاً من 4 مطورين لأكثر من 10 أشهر، مُدخلاً ممارسات Agile '
        'وتحسين اتساق وإنتاجية التسليم.')

    # Bayan Gardens
    company_line(doc, 'مدارس بيان \u2014 الخبر، المملكة العربية السعودية',
                 'ديسمبر 2017 \u2013 ديسمبر 2018')
    role_line(doc, 'مطور ويب وفني تقنية معلومات')
    bullet(doc,
        'طوّر منصة ويب شبيهة بأنظمة إدارة التعلم تُمكّن المعلمين من مشاركة '
        'المواد التعليمية مع أولياء الأمور والطلاب.')

    # ── SKILLS ────────────────────────────────────────────────────────────────

    section_header(doc, 'المهارات')

    skill_row(doc, 'الحوكمة',
              'أطر GRC، امتثال تقنية المعلومات، التحكم في الوصول، السياسات الأمنية')
    skill_row(doc, 'عمليات تقنية المعلومات',
              'إدارة الأجهزة، اشتراكات SaaS، إدارة الهوية، عمليات الانضمام والمغادرة')
    skill_row(doc, 'التسليم',
              'Agile / Scrum، PMI-ACP، حوكمة Sprint، التحكم في التغيير، Jira')
    skill_row(doc, 'الخلفية التقنية',
              'C#/.NET، Python، REST APIs، gRPC، React')
    skill_row(doc, 'السحابة والبنية التحتية',
              'Azure، Docker، Kubernetes (CKAD)، Microservices، Azure Service Bus')
    skill_row(doc, 'قواعد البيانات',
              'SQL Server، MongoDB')

    # ── EDUCATION & CERTIFICATIONS ────────────────────────────────────────────

    section_header(doc, 'التعليم والشهادات')

    edu_row(doc, 'بكالوريوس هندسة البرمجيات',
            'جامعة كونكورديا، مونتريال \u2014 أغسطس 2022  |  قائمة العميد 2021\u20132022')
    edu_row(doc, 'مطوّر تطبيقات Kubernetes المعتمد (CKAD)',
            'مؤسسة Linux \u2014 يوليو 2023')
    edu_row(doc, 'ممارس Agile المعتمد PMI-ACP',
            'معهد إدارة المشاريع \u2014 مارس 2022')

    # ── FOOTER ────────────────────────────────────────────────────────────────

    p = doc.add_paragraph()
    spacing(p, before=6, after=0)
    rtl_paragraph(p)
    bottom_border(p, color='e2e8f0', sz='4')
    r = p.add_run('السيرة الذاتية الكاملة والمشاريع والمراجع \u2190 aymanshehri.me')
    font(r, 8.5, color=MUTED)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Save ──────────────────────────────────────────────────────────────────

    output = 'Ayman_AlShehri_Resume_AR.docx'
    doc.save(output)
    print(f'Done: {output}')


if __name__ == '__main__':
    create_resume_ar()
