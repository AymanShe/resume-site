"""
generate_resume.py
Generates Ayman_AlShehri_Resume.docx — a one-page recruiter resume.

Usage:
    pip install python-docx
    python generate_resume.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Colors ───────────────────────────────────────────────────────────────────
TEAL  = RGBColor(0x0d, 0x94, 0x88)
NAVY  = RGBColor(0x0f, 0x17, 0x2a)
MUTED = RGBColor(0x64, 0x74, 0x8b)
BODY  = RGBColor(0x1e, 0x29, 0x3b)

# ── Layout ───────────────────────────────────────────────────────────────────
MARGIN        = 0.6   # inches
PAGE_WIDTH    = 8.5   # letter
CONTENT_WIDTH = PAGE_WIDTH - 2 * MARGIN  # 7.3 inches → used for right-tab stops


# ── Helpers ──────────────────────────────────────────────────────────────────

def font(run, size, bold=False, italic=False, color=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def spacing(paragraph, before=0, after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing = Pt(12)


def bottom_border(paragraph, color='0d9488', sz='4'):
    """Add a colored bottom border (used for section headers and the header rule)."""
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def right_tab(paragraph, inches=CONTENT_WIDTH):
    """Set a right-aligned tab stop at `inches` from the left margin."""
    pPr  = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(inches * 1440)))  # twips
    tabs.append(tab)
    pPr.append(tabs)


# ── Document builders ─────────────────────────────────────────────────────────

def section_header(doc, text):
    p = doc.add_paragraph()
    spacing(p, before=5, after=1)
    r = p.add_run(text)
    font(r, 10, bold=True, color=NAVY)
    bottom_border(p)
    return p


def company_line(doc, company_loc, dates, top=3):
    p = doc.add_paragraph()
    spacing(p, before=top, after=0)
    right_tab(p)
    r = p.add_run(company_loc)
    font(r, 10, bold=True, color=NAVY)
    r = p.add_run('\t' + dates)
    font(r, 9, color=MUTED)
    return p


def role_line(doc, role):
    p = doc.add_paragraph()
    spacing(p, before=0, after=1)
    r = p.add_run(role)
    font(r, 9.5, italic=True, color=TEAL)
    return p


def bullet(doc, text):
    p = doc.add_paragraph()
    spacing(p, before=1, after=0)
    pf = p.paragraph_format
    pf.left_indent        = Inches(0.18)
    pf.first_line_indent  = Inches(-0.14)
    r = p.add_run('\u2013 ')   # en-dash
    font(r, 9.5, color=MUTED)
    r = p.add_run(text)
    font(r, 9.5, color=BODY)
    return p


def skill_row(doc, label, value):
    p = doc.add_paragraph()
    spacing(p, before=2, after=0)
    r = p.add_run(label + ':  ')
    font(r, 9.5, bold=True, color=NAVY)
    r = p.add_run(value)
    font(r, 9.5, color=BODY)
    return p


def edu_row(doc, title, detail):
    p = doc.add_paragraph()
    spacing(p, before=2, after=0)
    r = p.add_run(title + '  ')
    font(r, 9.5, bold=True, color=NAVY)
    r = p.add_run(detail)
    font(r, 9.5, color=MUTED)
    return p


# ── Main ──────────────────────────────────────────────────────────────────────

def create_resume():
    doc = Document()

    # Page layout
    for sec in doc.sections:
        sec.page_width    = Inches(PAGE_WIDTH)
        sec.page_height   = Inches(11)
        sec.left_margin   = Inches(MARGIN)
        sec.right_margin  = Inches(MARGIN)
        sec.top_margin    = Inches(MARGIN)
        sec.bottom_margin = Inches(MARGIN)

    # Reset Normal style defaults
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(0)

    # ── HEADER ────────────────────────────────────────────────────────────────

    p = doc.add_paragraph()
    spacing(p, before=0, after=1)
    r = p.add_run('Ayman AlShehri')
    font(r, 22, bold=True, color=NAVY)

    p = doc.add_paragraph()
    spacing(p, before=0, after=2)
    r = p.add_run('IT Manager & Digital Excellence Leader  |  PMI-ACP  |  CKAD')
    font(r, 10, color=TEAL)

    p = doc.add_paragraph()
    spacing(p, before=0, after=4)
    r = p.add_run(
        'aymalsh@gmail.com  |  linkedin.com/in/aymanshe  |  '
        'aymanshehri.me  |  github.com/AymanShe  |  '
        '+966 550 832 888  |  Riyadh, Saudi Arabia'
    )
    font(r, 9, color=MUTED)
    bottom_border(p, color='0d9488', sz='6')

    # ── SUMMARY ───────────────────────────────────────────────────────────────

    section_header(doc, 'SUMMARY')

    p = doc.add_paragraph()
    spacing(p, before=3, after=0)
    r = p.add_run(
        'IT Manager with 7+ years across software engineering, technical leadership, and organizational '
        'governance. Currently building the operational infrastructure at RQM Business Solutions \u2014 '
        'GRC frameworks, IT operations standardization, and engineering delivery governance. '
        'PMI-ACP and CKAD certified; equally credible in a policy session and an architecture review.'
    )
    font(r, 9.5, color=BODY)

    # ── EXPERIENCE ────────────────────────────────────────────────────────────

    section_header(doc, 'EXPERIENCE')

    # RQM
    company_line(doc, 'RQM Business Solutions \u2014 Riyadh, Saudi Arabia', 'Oct 2024\u2013Present')
    role_line(doc, 'Digital Excellence Manager (Oct 2025\u2013Present)  /  Technical Lead (Oct 2024\u2013Oct 2025)')
    bullet(doc, 'Designed and enforced a GRC framework from scratch; standardized IT operations '
                'including device management, access control, SaaS subscriptions, and '
                'onboarding/offboarding processes.')
    bullet(doc, 'Established sprint governance and change control protocols \u2014 negotiated scope '
                'changes professionally with internal stakeholders and clients, protecting team '
                'capacity and delivery commitments.')
    bullet(doc, 'Owned architecture, technical roadmap, and code quality for a cloud-native SaaS '
                'platform across a team of 4\u20136 engineers.')

    # Canada College
    company_line(doc, 'Canada College \u2014 Montreal, Quebec', 'Jul 2022\u2013May 2025 (Part-time)')
    role_line(doc, 'Technical Instructor & Internship Advisor')
    bullet(doc, 'Designed and delivered a full database curriculum achieving 100% student pass rate; '
                'managed 50+ intern placements end-to-end with Montreal tech employers.')

    # Genetec
    company_line(doc, 'Genetec \u2014 Montreal, Quebec', 'Sep 2022\u2013Apr 2023')
    role_line(doc, 'Software Developer')
    bullet(doc, 'Shipped the media player feature end-to-end on Clearance (Genetec\u2019s digital '
                'evidence platform); resolved a cross-service data integrity bug causing corrupted '
                'case visibility for investigators.')

    # Ministry of Justice
    company_line(doc, 'Ministry of Justice \u2014 Riyadh, Saudi Arabia', 'Jan 2019\u2013Dec 2021')
    role_line(doc, 'Software Developer')
    bullet(doc, 'Owned the Social Cases module of Najiz, Saudi Arabia\u2019s national justice platform '
                'serving millions of citizens.')
    bullet(doc, 'Served as Tech Lead for a 4-developer team for 10+ months, introducing Agile practices '
                'and improving delivery consistency.')

    # Bayan Gardens
    company_line(doc, 'Bayan Gardens School \u2014 Khobar, Saudi Arabia', 'Dec 2017\u2013Dec 2018')
    role_line(doc, 'Web Developer & IT Technician')
    bullet(doc, 'Built an LMS-like web platform enabling teachers to digitally share materials with '
                'parents and students.')

    # ── SKILLS ────────────────────────────────────────────────────────────────

    section_header(doc, 'SKILLS')

    skill_row(doc, 'Governance',    'GRC frameworks, IT compliance, access control, security policy')
    skill_row(doc, 'IT Operations', 'Device management, SaaS subscriptions, identity management, onboarding/offboarding')
    skill_row(doc, 'Delivery',      'Agile / Scrum, PMI-ACP, sprint governance, change control, Jira')
    skill_row(doc, 'Backend',       'C#/.NET, Python, REST APIs, gRPC, React')
    skill_row(doc, 'Cloud & Infra', 'Azure, Docker, Kubernetes (CKAD), Microservices, Azure Service Bus')
    skill_row(doc, 'Data',          'SQL Server, MongoDB')

    # ── EDUCATION & CERTIFICATIONS ────────────────────────────────────────────

    section_header(doc, 'EDUCATION & CERTIFICATIONS')

    edu_row(doc, 'B.Eng. Software Engineering',
                 'Concordia University, Montreal \u2014 Aug 2022  |  Dean\u2019s List 2021\u20132022')
    edu_row(doc, 'Certified Kubernetes Application Developer (CKAD)',
                 'The Linux Foundation \u2014 Jul 2023')
    edu_row(doc, 'PMI Agile Certified Practitioner (PMI-ACP)',
                 'Project Management Institute \u2014 Mar 2022')

    # ── FOOTER ────────────────────────────────────────────────────────────────

    p = doc.add_paragraph()
    spacing(p, before=6, after=0)
    bottom_border(p, color='e2e8f0', sz='4')
    r = p.add_run('Full work history, projects, and references \u2192 aymanshehri.me')
    font(r, 8.5, italic=True, color=MUTED)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Save ──────────────────────────────────────────────────────────────────

    output = 'Ayman_AlShehri_Resume.docx'
    doc.save(output)
    print(f'Done: {output}')


if __name__ == '__main__':
    create_resume()
